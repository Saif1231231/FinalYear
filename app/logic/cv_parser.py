"""
CV PARSER (cv_parser.py)
------------------------
This file is responsible for "reading" the CV files that users upload.
It can read both PDF and Word (.docx) files.
"""

import io
import re
import logging

# We use this to keep track of any issues during reading
logger = logging.getLogger(__name__)

# These are common headings we look for to understand the structure of the CV
SECTION_HEADERS = {
    "SKILLS":      re.compile(r'\b(skills?|technical skills?|technologies|expertise|proficienc)\b', re.I),
    "EXPERIENCE":  re.compile(r'\b(experience|employment|work history|career|background)\b', re.I),
    "EDUCATION":   re.compile(r'\b(education|qualifications?|academic|degree|university)\b', re.I),
    "PROJECTS":    re.compile(r'\b(projects?|portfolio|personal projects?)\b', re.I),
    "CERTIFICATIONS": re.compile(r'\b(certifications?|certificates?|awards?)\b', re.I),
    "SUMMARY":     re.compile(r'\b(summary|profile|objective|about me|overview)\b', re.I),
}

def _clean_text(text):
    """Clean up the extracted text by removing weird characters and extra spaces."""
    # Remove non-standard characters
    cleaned = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\xA0-\uFFFF]', ' ', text)
    # Fix spacing issues
    cleaned = re.sub(r'[ \t]+', ' ', cleaned)
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    return cleaned.strip()

def _split_into_sections(text):
    """Try to figure out which part of the CV is the 'Skills' section, which is 'Experience', etc."""
    lines = text.split('\n')
    sections = {}
    current_section = 'OTHER' # Default section if we don't know where we are
    buffer = []

    for line in lines:
        clean_line = line.strip()
        found_header = None
        
        # Check if this line looks like a section heading
        for name, pattern in SECTION_HEADERS.items():
            if clean_line and len(clean_line) < 50 and pattern.search(clean_line):
                found_header = name
                break

        if found_header:
            # Save the text from the previous section
            if buffer:
                sections.setdefault(current_section, []).extend(buffer)
                buffer = []
            current_section = found_header
        else:
            buffer.append(line)

    # Save the last section
    if buffer:
        sections.setdefault(current_section, []).extend(buffer)

    # Combine the lines back into strings
    return {k: '\n'.join(v) for k, v in sections.items()}

def extract_text_from_pdf(file_bytes):
    """Read text from a PDF file."""
    text = ""
    try:
        # We try PyPDF2 first because it's fast
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            content = page.extract_text()
            if content:
                text += content + "\n"
    except Exception:
        # If PyPDF2 fails, we try another tool called fitz (PyMuPDF)
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
        except Exception as e:
            logger.error("Could not read PDF: %s", e)
    
    return _clean_text(text)

def extract_text_from_docx(file_bytes):
    """Read text from a Word (.docx) file."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        # Get text from paragraphs
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Get text from tables (many CVs use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text for cell in row.cells))
        
        return _clean_text("\n".join(parts))
    except Exception as e:
        logger.error("Could not read DOCX: %s", e)
        return ""

def extract_sections(file_bytes, filename):
    """Main function to get both the full text and the divided sections."""
    name_lower = filename.lower()
    
    if name_lower.endswith(".pdf"):
        raw_text = extract_text_from_pdf(file_bytes)
    elif name_lower.endswith((".docx", ".doc")):
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raw_text = ""

    sections = _split_into_sections(raw_text)
    return {"raw": raw_text, "sections": sections}

def is_allowed_file(filename):
    """Check if the file is a type we can read."""
    return filename.lower().endswith((".pdf", ".docx", ".doc"))

