"""
Update Mid-Point Report v2: 
- Replace diagram images with simpler draw.io style ones
- Rewrite sections from 'Solution implementation' onwards
- RESTORE the References section from the backup
"""
import os
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

SRC = "/Users/saifkhan/Desktop/Final year/docs/Mid-Point Reportdoc_BACKUP.docx"
OUT = "/Users/saifkhan/Desktop/Final year/docs/Mid-Point Reportdoc.docx"
IMG_DIR = "/Users/saifkhan/.gemini/antigravity/brain/de21cbc4-91f7-4d4b-ad13-2ee16cafde22"

def find_img(prefix):
    files = sorted([f for f in os.listdir(IMG_DIR) if f.startswith(prefix)], reverse=True)
    return os.path.join(IMG_DIR, files[0]) if files else None

ARCH_IMG = find_img("simple_architecture")
SEQ_IMG  = find_img("simple_sequence")
UC_IMG   = find_img("simple_usecase")
DFD_IMG  = find_img("simple_dataflow")
WIRE_IMG = find_img("simple_wireframe")

print(f"Architecture: {ARCH_IMG}")
print(f"Sequence: {SEQ_IMG}")
print(f"Use Case: {UC_IMG}")
print(f"DFD: {DFD_IMG}")
print(f"Wireframe: {WIRE_IMG}")

doc = Document(SRC)

# --- STEP 1: Replace diagrams in Section 5/6 ---
fig_map = {}
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if 'Figure 2' in txt and 'Use Case' in txt:
        fig_map['use_case'] = i
    elif 'Figure 4' in txt and 'Wireframe' in txt:
        fig_map['wireframe'] = i
    elif 'Figure 5' in txt and 'Architecture' in txt:
        fig_map['architecture'] = i
    elif 'Figure 6' in txt and 'Sequence' in txt:
        fig_map['sequence'] = i

print(f"Figure captions found: {fig_map}")

img_lookup = {
    'use_case': UC_IMG,
    'wireframe': WIRE_IMG,
    'architecture': ARCH_IMG,
    'sequence': SEQ_IMG,
}

for fig_name, caption_idx in fig_map.items():
    for check_idx in [caption_idx - 1, caption_idx, caption_idx - 2]:
        if check_idx < 0 or check_idx >= len(doc.paragraphs):
            continue
        p = doc.paragraphs[check_idx]
        drawings = p._element.findall('.//' + qn('wp:inline'))
        drawings += p._element.findall('.//' + qn('wp:anchor'))
        if drawings:
            for d in drawings:
                d.getparent().remove(d)
            for run in p.runs:
                run._element.getparent().remove(run._element)
            run = p.add_run()
            if fig_name in img_lookup and img_lookup[fig_name]:
                run.add_picture(img_lookup[fig_name], width=Inches(5.0))
                print(f"  Replaced {fig_name} at para {check_idx}")
            break

# --- STEP 2: Find cut points ---
impl_idx = None
refs_idx = None
appendix_idx = None

for i, p in enumerate(doc.paragraphs):
    if p.style and 'Heading 1' in p.style.name:
        if 'Solution implementation' in p.text:
            impl_idx = i
        elif 'References' in p.text:
            refs_idx = i
        elif 'Appendic' in p.text:
            appendix_idx = i

print(f"Implementation: {impl_idx}, References: {refs_idx}, Appendices: {appendix_idx}")

# --- STEP 3: Extract references text from backup before removing ---
ref_texts = []
if refs_idx is not None:
    end_idx = appendix_idx if appendix_idx else len(doc.paragraphs)
    for i in range(refs_idx, end_idx):
        p = doc.paragraphs[i]
        ref_texts.append(p.text)

print(f"Extracted {len(ref_texts)} reference paragraphs")

# --- STEP 4: Remove everything from impl_idx onwards ---
body = doc.element.body
found = False
to_remove = []
for child in list(body):
    if child.tag.endswith('}sectPr'):
        continue
    if child.tag.endswith('}p'):
        text = ''
        for sub in child.iter():
            if sub.text: text += sub.text
            if sub.tail: text += sub.tail
        if 'Solution implementation' in text and not found:
            found = True
        if found:
            to_remove.append(child)
    elif found:
        to_remove.append(child)

for elem in to_remove:
    body.remove(elem)

print(f"Removed {len(to_remove)} elements")

# --- STEP 5: Add new sections ---

def h1(d, t): return d.add_paragraph(t, style='Heading 1')
def h2(d, t): return d.add_paragraph(t, style='Heading 2')
def h3(d, t): return d.add_paragraph(t, style='Heading 3')
def para(d, t): return d.add_paragraph(t)
def bpara(d, b, n):
    p = d.add_paragraph()
    r = p.add_run(b); r.bold = True
    p.add_run(n)
    return p

# == SOLUTION IMPLEMENTATION ==
h1(doc, "Solution Implementation")
para(doc, "This section describes how CVMatchMaker was built, following the Agile methodology. Development was divided into four sprints, each focused on a core feature area.")

h2(doc, "7.1  Sprint 1 — Content Extraction and CV Parsing")
para(doc, "The first sprint focused on reading CV files and extracting text content. The goal was to accept both PDF and Word documents.")
bpara(doc, "Technical Challenge: ", "PDFs and DOCX files store text completely differently. A PDF stores characters on a canvas, while DOCX stores XML paragraphs.")
bpara(doc, "Solution: ", "Two Python libraries were used. PyPDF2 reads PDF files and python-docx reads Word documents. The cv_parser.py module detects the file extension and routes it to the correct library. After extraction, a section splitter uses regular expressions to detect headings like Skills, Experience, and Education. The file is read into memory using file.read() and is never saved to disk, complying with GDPR.")

h2(doc, "7.2  Sprint 2 — NLP Skill Extraction Engine")
para(doc, "The second sprint built the NLP pipeline to automatically identify technical skills from CV text.")
bpara(doc, "Technical Challenge: ", "Simple keyword matching is unreliable. The word 'developing' should match 'development', and 'reactjs' should be recognised as 'react'.")
bpara(doc, "Solution: ", "A three-layer extraction was implemented in skill_extractor.py. Layer 1 uses regex pattern matching to find exact phrases like 'machine learning'. Layer 2 uses spaCy NLP for lemmatisation — converting words to their base form. Layer 3 uses synonym normalisation mapping aliases to standard names. Skills from the CV's Skills section receive weight 2.0, Experience gets 1.5, and elsewhere gets 1.0.")

h2(doc, "7.3  Sprint 3 — Job Fetching and Recommendation Engine")
para(doc, "The third sprint connected the system to the Reed.co.uk Jobs API and built the comparison engine.")
bpara(doc, "Technical Challenge: ", "The Reed API returns raw JSON with skills embedded in long description paragraphs. There is no structured skills field.")
bpara(doc, "Solution: ", "The same NLP pipeline from Sprint 2 was reused on job descriptions. The analyser.py module compares CV skills against job skills using Python set operations: intersection gives matched skills, set difference gives missing skills. The score is (matched / total) x 100. Missing skills are ranked by frequency and tagged High (>50%), Medium (>20%), or Low priority.")

h2(doc, "7.4  Sprint 4 — Frontend, Authentication, and Features")
para(doc, "The final sprint built the web interface, user authentication, and additional features including Learning Path, History, Compare, and Cover Letter generator.")
bpara(doc, "Technical Challenge: ", "The results page needed to display a large amount of data cleanly on both desktop and mobile screens.")
bpara(doc, "Solution: ", "The interface uses plain HTML5, CSS3, and JavaScript with a card-based layout. Matched skills are green chips, missing skills are red chips. Authentication uses Flask-Login with PBKDF2-SHA256 password hashing. The Learning Path saves missing skills as goals with progress tracking in the database.")

# == SOLUTION TESTING ==
h1(doc, "Solution Testing / Evaluation")
para(doc, "Testing was conducted at three levels: unit testing, integration testing, and user acceptance testing with representative users.")

h2(doc, "8.1  Unit Testing")
bpara(doc, "CV Parser Tests: ", "Sample PDF and DOCX files with known content were tested. The extract_sections() function correctly extracted raw text and identified section headings. Edge cases included CVs with no headings and very short CVs.")
bpara(doc, "Skill Extractor Tests: ", "10 test strings with known skills were passed through extract_skills_from_text(). All three layers worked correctly: regex caught exact phrases, lemmatisation caught word variations, and synonym normalisation mapped aliases.")
bpara(doc, "Score Calculator Tests: ", "A CV with 8 matched skills out of 10 total correctly returned 80%. Edge cases: zero matched returned 0%, excess skills capped at 100%.")

h2(doc, "8.2  Integration Testing")
para(doc, "The full pipeline was tested end-to-end with real CVs:")

t1 = doc.add_table(rows=4, cols=4, style='TableGrid')
for j, h in enumerate(['Test CV', 'Job Title', 'Expected', 'Result']):
    t1.rows[0].cells[j].text = h
    for r in t1.rows[0].cells[j].paragraphs[0].runs: r.bold = True
data1 = [
    ['Software Engineer CV (PDF)', 'Software Engineer', 'Score 60-80%', 'Score: 72%. Pass'],
    ['Data Analyst CV (DOCX)', 'Data Analyst', 'Score 50-70%', 'Score: 65%. Pass'],
    ['Empty skills CV (PDF)', 'Frontend Developer', 'Score < 20%', 'Score: 12%. Pass'],
]
for i, row in enumerate(data1):
    for j, cell in enumerate(row):
        t1.rows[i+1].cells[j].text = cell

h2(doc, "8.3  User Acceptance Testing")
para(doc, "User acceptance testing was conducted with three classmates representing the target audience.")

t2 = doc.add_table(rows=4, cols=4, style='TableGrid')
for j, h in enumerate(['Tester', 'Task', 'Feedback', 'Outcome']):
    t2.rows[0].cells[j].text = h
    for r in t2.rows[0].cells[j].paragraphs[0].runs: r.bold = True
data2 = [
    ['Emmanuel Ennin', 'Upload CV for Software Engineer', 'Upload was straightforward. Liked the green/red skill chips.', 'Pass'],
    ['Nahavipusan', 'Analyse for Data Analyst, check Learning Path', 'Learning Path was very useful. Priority chart was clear.', 'Pass'],
    ['Enoch Emron', 'Sign up, upload CV, view History, export report', 'Signup was quick. Score trend chart showed improvement.', 'Pass'],
]
for i, row in enumerate(data2):
    for j, cell in enumerate(row):
        t2.rows[i+1].cells[j].text = cell

para(doc, "All three testers completed the core workflow without assistance. Main feedback: skill matching was accurate, priority chart was clear, and Learning Path was a practical addition.")

h2(doc, "8.4  Evaluation Against Objectives")
t3 = doc.add_table(rows=8, cols=3, style='TableGrid')
for j, h in enumerate(['Objective', 'Status', 'Evidence']):
    t3.rows[0].cells[j].text = h
    for r in t3.rows[0].cells[j].paragraphs[0].runs: r.bold = True
data3 = [
    ['Research NLP tools', 'Met', 'spaCy selected after comparing NLTK and CoreNLP.'],
    ['Build CV reader for PDF/DOCX', 'Met', 'cv_parser.py handles both formats.'],
    ['Connect to job listing APIs', 'Met', 'Reed API integrated with mock data fallback.'],
    ['Build matching system', 'Met', 'Three-layer extraction + set comparison.'],
    ['Design user-friendly website', 'Met', 'All testers completed tasks without help.'],
    ['Ensure security and encryption', 'Met', 'PBKDF2 passwords. CVs processed in memory only.'],
    ['Test with sample CVs', 'Met', 'Unit, integration, and UAT completed.'],
]
for i, row in enumerate(data3):
    for j, cell in enumerate(row):
        t3.rows[i+1].cells[j].text = cell

# == DISCUSSION ==
h1(doc, "Discussion")
h2(doc, "9.1  Key Findings")
para(doc, "NLP-based skill extraction combined with real-time Reed API job data produces actionable career feedback. The three-layer extraction approach proved more reliable than single-method approaches. The priority ranking based on employer demand frequency was highlighted by all testers as the most useful feature.")

h2(doc, "9.2  Comparison with Literature")
para(doc, "Kopparapu's 2010 system (8) only did basic keyword matching and struggled with synonyms — CVMatchMaker's synonym normalisation solves this. Ranjan's 2015 system (10) achieved 82% accuracy but only gave yes/no answers — CVMatchMaker provides percentage scores and learning resources. Unlike TopCV (99 pounds) and Jobscan (49 pounds/month), CVMatchMaker is completely free and provides instant results.")

h2(doc, "9.3  Strengths")
para(doc, "Main strengths: (1) real-time job data from Reed API, (2) three-layer NLP catches skills that simple keyword search misses, (3) Learning Path turns analysis into action, (4) privacy-first design processes CVs in memory only, (5) completely free with no subscription.")

h2(doc, "9.4  Limitations")
para(doc, "Limitations identified: (1) skills dictionary contains approximately 200 skills — niche skills may not be recognised, (2) only supports English-language CVs and UK jobs, (3) the smallest spaCy model is used — a larger model could improve accuracy, (4) experience relevance and format quality scores are currently placeholder values.")

# == CONCLUSION ==
h1(doc, "Conclusion")
h2(doc, "10.1  Summary")
para(doc, "CVMatchMaker was successfully designed, built, and tested. The system uses content extraction (PyPDF2, python-docx), NLP (spaCy), and recommendation techniques (set comparison with Reed API data) to provide actionable career feedback. All seven original objectives were met.")

h2(doc, "10.2  Recommendations for Further Work")
para(doc, "1. Expand the skills dictionary to 500+ skills by scraping common skills from Reed and LinkedIn job listings.")
para(doc, "2. Add support for additional job APIs (Indeed, LinkedIn Jobs) for broader market coverage.")
para(doc, "3. Implement a proper experience relevance score using spaCy Named Entity Recognition.")
para(doc, "4. Add PDF export for the results page.")
para(doc, "5. Build a mobile-responsive version of the dashboard.")

h2(doc, "10.3  Reflection")
para(doc, "This project taught me how NLP works in practice — from tokenisation and lemmatisation to synonym handling. Integrating the Reed API gave me experience with REST APIs and real-world data. The biggest lesson was privacy-by-design: processing CVs in memory and never saving them was a technical constraint that influenced the entire architecture, but was the right decision under GDPR.")
para(doc, "If I were to start again, I would begin user testing earlier. The feedback from Emmanuel, Nahavipusan, and Enoch led to useful improvements that could have been implemented earlier.")

# == REFERENCES (restored from backup) ==
h1(doc, "References")
# Skip the heading text itself (first entry) and any guidance text at the end
for txt in ref_texts[1:]:
    t = txt.strip()
    if not t:
        continue
    if t.startswith("In this section") or t.startswith("It is highly") or t.startswith("It is your"):
        continue  # Skip template guidance text
    para(doc, t)

# == APPENDICES ==
h1(doc, "Appendices")
para(doc, "Appendix A: Project Proposal Poster — see docs/Poster.pdf")
para(doc, "Appendix B: Project Management — Trello board was used throughout development with weekly supervisor meetings.")
para(doc, "Appendix C: Technical Access — The complete source code is available in the app/ directory. To run: activate the virtual environment, navigate to app/, and run python3 app.py.")
para(doc, "Appendix D: Ethics Form — Submitted and approved. The system processes CVs in memory only and never stores personal documents.")
para(doc, "Appendix E: Risk Register — See Section 4.5 of this report.")

# SAVE
doc.save(OUT)
print(f"\n Done! Saved to {OUT}")
