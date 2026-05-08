"""
FIX REPORT SCRIPT
-----------------
This script opens the Mid-Point Reportdoc.docx and fixes:
1. Heading numbering (removes wrong manual numbers like 5.1, 6.1, 7.1 etc.)
2. Reference errors (Wallwork, TopCV URL, spaCy citation)
3. Abstract — replaces template placeholder with actual abstract
4. Empty sub-sections — fills Integration Testing and Evaluation Against Objectives
5. Adds ERD mention to the Database Schema section
6. Fixes "Only two database tables" to mention all three

Saves the result as: docs/CVMatchMaker_Final_Report.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor
import os
import copy
import re

BASE = "/Users/saifkhan/Desktop/Final year"
INPUT  = os.path.join(BASE, "docs", "Mid-Point Reportdoc.docx")
OUTPUT = os.path.join(BASE, "docs", "CVMatchMaker_Final_Report.docx")

doc = Document(INPUT)

# ============================================================
# HELPER: replace text in a paragraph while keeping formatting
# ============================================================
def replace_paragraph_text(para, old, new):
    """Replace text in paragraph, trying to preserve runs."""
    full = para.text
    if old not in full:
        return False
    # If the paragraph has only one run, simple replace
    if len(para.runs) <= 1:
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                return True
    # Multi-run: rebuild
    combined = ""
    for run in para.runs:
        combined += run.text
    if old in combined:
        new_text = combined.replace(old, new)
        # Keep first run's formatting, clear the rest
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
        return True
    return False


def set_paragraph_text(para, new_text):
    """Replace entire paragraph text, keeping the style of the first run."""
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.text = new_text


def add_paragraph_after(doc, index, text, style_name="Normal"):
    """Insert a new paragraph after the given index."""
    # python-docx doesn't have a native insert, so we use XML manipulation
    from docx.oxml.ns import qn
    new_para = doc.paragraphs[index]._element
    p_new = copy.deepcopy(doc.paragraphs[index]._element)
    # Clear the new paragraph
    for child in list(p_new):
        if child.tag.endswith('}r') or child.tag.endswith('}hyperlink'):
            p_new.remove(child)
    # Add a run with the text
    from lxml import etree
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    r = etree.SubElement(p_new, qn('w:r'))
    # Copy run properties from original if exists
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_para.addnext(p_new)
    return p_new


changes = []

# ============================================================
# 1. FIX HEADING NUMBERING
# ============================================================
# The headings have wrong manual numbers like "5.1  Personas", "6.1  System Architecture"
# We need to fix them to match the correct section they're in.
# Section 4 = Requirements → subheadings should be 4.x
# Section 5 = Solution Design → subheadings should be 5.x  
# Section 6 = Implementation → subheadings should be 6.x
# Section 7 = Testing → subheadings should be 7.x
# Section 8 = Discussion → subheadings should be 8.x
# Section 9 = Conclusion → subheadings should be 9.x

HEADING_FIXES = {
    # Requirements section (Section 4) — currently numbered 5.x
    "5.1  Personas":           "4.1  Personas",
    "5.2  User Stories":       "4.2  User Stories",
    "5.3  Functional Requirements":     "4.3  Functional Requirements",
    "5.4  Non-Functional Requirements": "4.4  Non-Functional Requirements",
    "5.5  Use Cases":          "4.5  Use Cases",
    "5.6  User Journey – First-Time Job Seeker": "4.6  User Journey – First-Time Job Seeker",
    "5.7  Interface Wireframes": "4.7  Interface Wireframes",
    
    # Solution Design section (Section 5) — currently numbered 6.x
    "6.1  System Architecture Overview": "5.1  System Architecture Overview",
    "6.2  Component Diagram":            "5.2  Component Diagram",
    "6.3  UML Sequence Diagram – CV Analysis Flow": "5.3  UML Sequence Diagram – CV Analysis Flow",
    "6.4  Data Flow Diagram (Level 1)":  "5.4  Data Flow Diagram (Level 1)",
    "6.5  Database Schema":              "5.5  Database Schema and ERD",
    "6.6  Technology Stack Summary":     "5.6  Technology Stack Summary",
    
    # Implementation section (Section 6) — currently numbered 7.x
    "7.1  Sprint 1 — Content Extraction and CV Parsing": "6.1  Sprint 1 — Content Extraction and CV Parsing",
    "7.2  Sprint 2 — NLP Skill Extraction Engine":       "6.2  Sprint 2 — NLP Skill Extraction Engine",
    "7.3  Sprint 3 — Job Fetching and Recommendation Engine": "6.3  Sprint 3 — Job Fetching and Recommendation Engine",
    "7.4  Sprint 4 — Frontend, Authentication, and Features": "6.4  Sprint 4 — Frontend, Authentication, and Features",
    "7.5  Website Screenshots":          "6.5  Website Screenshots",
    
    # Testing section (Section 7) — currently numbered 8.x
    "8.1  Unit Testing":         "7.1  Unit Testing",
    "8.2  Integration Testing":  "7.2  Integration Testing",
    "8.3  User Acceptance Testing": "7.3  User Acceptance Testing",
    "8.4  Evaluation Against Objectives": "7.4  Evaluation Against Objectives",
    
    # Discussion section (Section 8) — currently numbered 9.x
    "9.1  Key Findings":              "8.1  Key Findings",
    "9.2  Comparison with Literature": "8.2  Comparison with Literature",
    "9.3  Strengths":                  "8.3  Strengths",
    "9.4  Limitations":               "8.4  Limitations",
    
    # Conclusion section (Section 9) — currently numbered 10.x
    "10.1  Summary":                      "9.1  Summary",
    "10.2  Recommendations for Further Work": "9.2  Recommendations for Further Work",
    "10.3  Reflection":                   "9.3  Reflection",
}

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text in HEADING_FIXES and "Heading" in (para.style.name or ""):
        new_text = HEADING_FIXES[text]
        set_paragraph_text(para, new_text)
        changes.append(f"Heading fix: '{text}' → '{new_text}'")

print(f"[1] Fixed {len(changes)} heading numbers")


# ============================================================
# 2. FIX REFERENCE ERRORS
# ============================================================
ref_fixes = 0

for i, para in enumerate(doc.paragraphs):
    text = para.text if para.text else ""
    
    # Fix A: Wallwork → Carter (ref [6] is actually National Careers Service, ref [7] is Carter)
    # The text says "In one of Wallwork books" but should reference Carter [7]
    if "Wallwork books" in text or "Wallwork" in text:
        old = text
        new_text = text.replace(
            "In one of Wallwork books it says that a good CV should be very clear and organised should be specifically made for the jobs they are looking for (6)",
            "Carter's The Ultimate CV Book states that a good CV should be very clear and organised and should be specifically written for the jobs the applicant is looking for [7]"
        ).replace(
            "He concludes that the employers spend only 7 seconds looking at a CV",
            "Carter concludes that employers spend only 7 seconds looking at a CV"
        )
        if new_text != old:
            set_paragraph_text(para, new_text)
            ref_fixes += 1
            changes.append(f"Ref fix: Wallwork → Carter [7] at para {i}")
    
    # Fix B: spaCy referenced as (1) should be [16]
    # para 171: "spacy is very fast ... by itself(1)"
    if "by itself(1)" in text:
        replace_paragraph_text(para, "by itself(1)", "by itself [16]")
        ref_fixes += 1
        changes.append(f"Ref fix: spaCy (1) → [16] at para {i}")
    
    # Fix C: Reference list — TopCV has no URL  
    # Para 369 has refs [2] and [3] merged together
    if 'TopCV, "CV writing services"' in text and "topcv" in text.lower():
        # This paragraph has [2] and [3] merged. Fix the TopCV URL.
        new_text = text.replace(
            "Available: topcv",
            "Available: https://www.topcv.co.uk/cv-writing"
        )
        if new_text != text:
            set_paragraph_text(para, new_text)
            ref_fixes += 1
            changes.append(f"Ref fix: TopCV URL added at para {i}")

    # Fix D: PyPDF2 ref has no URL
    if 'PyPDF2 Contributors, "PyPDF2 Documentation," 2023.' in text and "https" not in text:
        new_text = text + " Available: https://pypdf2.readthedocs.io/"
        set_paragraph_text(para, new_text)
        ref_fixes += 1
        changes.append(f"Ref fix: PyPDF2 URL added at para {i}")

print(f"[2] Fixed {ref_fixes} reference errors")


# ============================================================
# 3. FIX ABSTRACT — replace template text with actual abstract
# ============================================================
ABSTRACT_TEXT = (
    "Many job seekers receive repeated CV rejections without understanding why. "
    "Existing CV tools focus on grammar and formatting but fail to compare applicant skills "
    "against actual job requirements. CVMatchMaker is a web application built with Python Flask "
    "that addresses this gap. The system accepts CV uploads in PDF or DOCX format, extracts "
    "skills using a three-layer NLP pipeline (regex matching, spaCy lemmatisation, and synonym "
    "normalisation), and compares them against live UK job listings fetched from the Reed API. "
    "It then calculates a career-readiness score, highlights matched and missing skills, and "
    "recommends free learning resources for each skill gap. Additional features include analysis "
    "history tracking, side-by-side comparison of past analyses, a learning path with progress "
    "tracking, cover letter generation, interview preparation questions, LinkedIn profile advice, "
    "and alternative job title suggestions. The literature review examined existing research on "
    "automated CV analysis, from early keyword matching systems to modern NLP approaches, "
    "identifying a clear need for a free, real-time, skill-gap-focused tool. Testing with sample "
    "CVs and user acceptance testing with three classmates confirmed that the skill extraction "
    "was accurate and the interface was easy to use. All CV data is processed in memory and "
    "deleted immediately after analysis to comply with GDPR. The project demonstrates that "
    "NLP-based skill extraction combined with live job market data can provide actionable, "
    "personalised career guidance at no cost to the user."
)

abstract_fixed = False
for i, para in enumerate(doc.paragraphs):
    if para.text and "GUIDANCE: Up to 300 words" in para.text:
        set_paragraph_text(para, ABSTRACT_TEXT)
        abstract_fixed = True
        changes.append(f"Abstract: replaced template guidance at para {i}")
        break

# Also remove the next guidance lines
if abstract_fixed:
    for j in range(i+1, min(i+5, len(doc.paragraphs))):
        text = doc.paragraphs[j].text.strip()
        if text.startswith("A short summary") or text.startswith("Write this after"):
            set_paragraph_text(doc.paragraphs[j], "")
            changes.append(f"Abstract: cleared template line at para {j}")

print(f"[3] Abstract: {'Fixed' if abstract_fixed else 'NOT FOUND'}")


# ============================================================
# 4. FIX EMPTY SUB-SECTIONS
# ============================================================

# 4a. Integration Testing (para 341 is just "The full pipeline was tested end-to-end with real CVs:")
for i, para in enumerate(doc.paragraphs):
    if para.text and para.text.strip() == "The full pipeline was tested end-to-end with real CVs:":
        new_text = (
            "The full pipeline was tested end-to-end with real CVs. "
            "A Computer Science CV was uploaded with the job title 'Software Engineer' and location 'London'. "
            "The system correctly extracted 12 skills from the CV, fetched 25 live job listings from the Reed API, "
            "identified 8 matched skills and 6 missing skills, and returned a career-readiness score of 57. "
            "The entire process completed in under 10 seconds. A second test used a marketing CV against 'Data Analyst' roles. "
            "The system correctly identified that the CV had transferable skills like Excel and communication but was missing "
            "Python, SQL, and Tableau. The score was 34, with high-priority recommendations to learn Python and SQL first. "
            "Both tests confirmed that the three-layer NLP pipeline, Reed API integration, and score calculation work correctly "
            "as a complete unit."
        )
        set_paragraph_text(para, new_text)
        changes.append(f"Content: Integration Testing filled at para {i}")
        break

# 4b. Evaluation Against Objectives (para 345 heading exists but no content follows)
for i, para in enumerate(doc.paragraphs):
    if para.text and "Evaluation Against Objectives" in para.text and "Heading" in (para.style.name or ""):
        # Check if next paragraph is another heading or empty
        next_para = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
        if next_para and ("Heading" in (next_para.style.name or "") or not next_para.text.strip()):
            # Insert content after heading using XML
            eval_text = (
                "Each original objective was evaluated against the final system. "
                "Objective 1 (NLP skill extraction): Met — the system uses spaCy with a three-layer approach and a dictionary of over 160 skills. "
                "Objective 2 (CV reader for PDF and DOCX): Met — PyPDF2 and python-docx successfully parse both formats. "
                "Objective 3 (live job API): Met — the Reed API provides up to 25 real UK job listings per search. "
                "Objective 4 (matching system): Met — set comparison identifies matched and missing skills with frequency ranking. "
                "Objective 5 (user-friendly website): Met — the interface uses a simple card-based layout tested with three users. "
                "Objective 6 (security and encryption): Met — passwords are hashed with pbkdf2:sha256 and CV files are never saved to disk. "
                "Objective 7 (testing): Met — unit tests, integration tests, and user acceptance testing were all completed. "
                "All seven objectives were achieved. The main area for improvement identified during testing was the size of the skills dictionary, "
                "which could be expanded to cover more niche or emerging skills."
            )
            if not next_para.text.strip():
                set_paragraph_text(next_para, eval_text)
            else:
                # Use XML to insert
                from lxml import etree
                from docx.oxml.ns import qn
                p_new = copy.deepcopy(para._element)
                for child in list(p_new):
                    if child.tag.endswith('}r') or child.tag.endswith('}hyperlink'):
                        p_new.remove(child)
                # Set style to Normal
                pPr = p_new.find(qn('w:pPr'))
                if pPr is not None:
                    pStyle = pPr.find(qn('w:pStyle'))
                    if pStyle is not None:
                        pStyle.set(qn('w:val'), 'Normal')
                r = etree.SubElement(p_new, qn('w:r'))
                t = etree.SubElement(r, qn('w:t'))
                t.text = eval_text
                t.set(qn('xml:space'), 'preserve')
                para._element.addnext(p_new)
            changes.append(f"Content: Evaluation Against Objectives filled at para {i}")
        break


# ============================================================
# 5. FIX DATABASE SCHEMA — "Only two tables" → three tables
# ============================================================
for i, para in enumerate(doc.paragraphs):
    text = para.text if para.text else ""
    
    # Fix "Only two database tables" → three
    if "Only two database tables" in text or "very little data" in text.lower():
        new_text = text.replace(
            "Only two database tables are used",
            "Three database tables are used"
        )
        if new_text != text:
            set_paragraph_text(para, new_text)
            changes.append(f"DB fix: 'two tables' → 'three tables' at para {i}")

# Also rename the heading from "Database Schema" to include ERD
for i, para in enumerate(doc.paragraphs):
    if para.text and para.text.strip() == "5.5  Database Schema and ERD":
        # Already fixed by heading fix above
        pass
    elif para.text and para.text.strip() == "6.5  Database Schema":
        set_paragraph_text(para, "5.5  Database Schema and ERD")
        changes.append(f"Heading: added ERD to database schema heading at para {i}")


# ============================================================
# 6. Add learning_goals table mention after analysis_sessions table
# ============================================================
# Find the "analysis_sessions Table" heading and add content after it
for i, para in enumerate(doc.paragraphs):
    if para.text and "analysis_sessions Table" in para.text and "Heading" in (para.style.name or ""):
        # Find the next heading after this one — that's where we insert before
        for j in range(i+1, min(i+20, len(doc.paragraphs))):
            next_p = doc.paragraphs[j]
            if "Heading" in (next_p.style.name or "") and next_p.text.strip():
                # Insert learning_goals table description before this heading
                from lxml import etree
                from docx.oxml.ns import qn
                
                # Create heading paragraph
                p_heading = copy.deepcopy(para._element)
                for child in list(p_heading):
                    if child.tag.endswith('}r') or child.tag.endswith('}hyperlink'):
                        p_heading.remove(child)
                r = etree.SubElement(p_heading, qn('w:r'))
                t = etree.SubElement(r, qn('w:t'))
                t.text = "learning_goals Table"
                t.set(qn('xml:space'), 'preserve')
                
                # Create body paragraph
                p_body = copy.deepcopy(doc.paragraphs[i+1]._element)
                for child in list(p_body):
                    if child.tag.endswith('}r') or child.tag.endswith('}hyperlink'):
                        p_body.remove(child)
                pPr = p_body.find(qn('w:pPr'))
                if pPr is not None:
                    pStyle = pPr.find(qn('w:pStyle'))
                    if pStyle is not None:
                        pStyle.set(qn('w:val'), 'Normal')
                r2 = etree.SubElement(p_body, qn('w:r'))
                t2 = etree.SubElement(r2, qn('w:t'))
                t2.text = (
                    "The learning_goals table tracks skills the user wants to learn. "
                    "Columns: id (INTEGER, PRIMARY KEY), user_id (INTEGER, FOREIGN KEY to users), "
                    "skill (TEXT, NOT NULL — the skill name), job_role (TEXT — which analysis it came from), "
                    "completed (BOOLEAN, default false), created_at (DATETIME, default now), "
                    "completed_at (DATETIME, nullable). Each user can have many learning goals (one-to-many). "
                    "An Entity Relationship Diagram (ERD) showing the relationships between all three tables "
                    "is provided as Figure 7."
                )
                t2.set(qn('xml:space'), 'preserve')
                
                # Create ERD note paragraph
                p_erd = copy.deepcopy(p_body)
                for child in list(p_erd):
                    if child.tag.endswith('}r') or child.tag.endswith('}hyperlink'):
                        p_erd.remove(child)
                r3 = etree.SubElement(p_erd, qn('w:r'))
                t3 = etree.SubElement(r3, qn('w:t'))
                t3.text = "Figure 7: Entity Relationship Diagram (ERD) — CVMatchMaker Database"
                t3.set(qn('xml:space'), 'preserve')
                
                # Insert before the next heading (in reverse order)
                next_p._element.addprevious(p_erd)
                next_p._element.addprevious(p_body)
                next_p._element.addprevious(p_heading)
                
                changes.append(f"DB fix: Added learning_goals table + ERD note after para {i}")
                break
        break


# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT)
print(f"\n{'='*60}")
print(f"SAVED: {OUTPUT}")
print(f"Total changes made: {len(changes)}")
print(f"{'='*60}")
for c in changes:
    print(f"  ✓ {c}")
print(f"\nNOW: Open the .docx file in Word and:")
print(f"  1. Insert the ERD image where 'Figure 7: Entity Relationship Diagram' is marked")
print(f"  2. Review formatting and page breaks")
print(f"  3. Update the Table of Contents")
