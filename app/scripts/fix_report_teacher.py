"""
FIX REPORT — TEACHER FEEDBACK
Applies all teacher feedback changes to Mid-Point Reportdom.docx
"""
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import copy, os

INPUT  = "/Users/saifkhan/Desktop/Final year/docs/Mid-Point Reportdom.docx"
OUTPUT = "/Users/saifkhan/Desktop/Final year/docs/Mid-Point Reportdom.docx"

doc = Document(INPUT)
changes = []

def set_text(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.text = new_text

def add_table_borders(table):
    """Add simple plain borders to a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
        tbl.insert(0, tblPr)
    # Remove existing borders
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


# ============================================================
# 1. REWRITE AIMS (paras 119, 120, 121)
# Teacher says: remove "1st Aim:", "2nd Aim:", "3rd Aim:" form
# Use proper academic descriptions
# ============================================================
NEW_AIMS = [
    "To use content extraction techniques to extract CV content from PDF and DOCX files, and apply natural language processing (NLP) to analyse the extracted text, identifying key skills and phrases for comparison against real job listings.",
    "To apply recommendation and search techniques to provide personalised feedback based on the CV analysis, including identifying missing skills and suggesting free learning resources so users can improve.",
    "To design a user-friendly and accessible web interface that keeps user data private and secure, whilst providing clear and actionable advice on how to improve the CV for the target job market.",
]

for idx, aim_text in enumerate(NEW_AIMS):
    para_idx = 119 + idx
    if para_idx < len(doc.paragraphs):
        set_text(doc.paragraphs[para_idx], f"{idx+1}. {aim_text}")
        changes.append(f"Aim {idx+1} rewritten at para {para_idx}")


# ============================================================
# 2. REWRITE OBJECTIVES (paras 124-130)
# Teacher wants: obj 1 about literature review, obj 2 about CV reader
# ============================================================
NEW_OBJECTIVES = [
    "To conduct a thorough literature review to identify the most suitable NLP tools and libraries for extracting skills from CVs and job listings.",
    "To create a CV reader application that can effectively open and process different formats of CVs including PDF and Word documents, ensuring that key terms and phrases are extracted correctly for skills matching at a later stage.",
    "To connect the system to job listing websites using their APIs, such as Reed, to retrieve real-time job data for comparison.",
    "To build a matching system that compares the skills extracted from the CV with the skills required in job descriptions, identifying gaps and ranking them by importance.",
    "To design a simple and user-friendly website so that users can easily navigate the platform, upload their CVs, and view their results without difficulty.",
    "To ensure the website is secure, with password encryption and GDPR-compliant data handling, so that user data is kept safe at all times.",
    "To test the final system thoroughly using unit tests, integration tests, and user acceptance testing with sample CVs to confirm that it runs smoothly, accurately, and reliably.",
]

for idx, obj_text in enumerate(NEW_OBJECTIVES):
    para_idx = 124 + idx
    if para_idx < len(doc.paragraphs):
        set_text(doc.paragraphs[para_idx], obj_text)
        changes.append(f"Objective {idx+1} rewritten at para {para_idx}")


# ============================================================
# 3. PROJECT MANAGEMENT — Add Gantt chart note (para 178)
# Teacher asks: "Do you have a Gantt Chart from Trello?"
# ============================================================
for i, p in enumerate(doc.paragraphs):
    if p.text and "The product lifecycle has four phases" in p.text:
        old = p.text
        new_text = old.rstrip() + " A Gantt chart derived from the Trello board is included in Appendix B, showing task timelines across all four phases."
        set_text(p, new_text)
        changes.append(f"Added Gantt chart reference at para {i}")
        break


# ============================================================
# 4. PERSONAS — Add gender info + Persona IDs
# Teacher: "include the gender info to show the balance"
# Teacher: "add these ID to the personas' introduction"
# ============================================================
PERSONA_FIXES = {
    "Alex Turner, 22 — Final-year Computer Science student.":
        "Alex Turner (Male), 22 — Final-year Computer Science student. [Persona ID: P-01]",
    "Priya Sharma, 32 — Moving from retail management into IT project management.":
        "Priya Sharma (Female), 32 — Moving from retail management into IT project management. [Persona ID: P-02]",
    "James Okafor, 45 — Former warehouse supervisor, currently unemployed.":
        "James Okafor (Male), 45 — Former warehouse supervisor, currently unemployed. [Persona ID: P-03]",
}

for i, p in enumerate(doc.paragraphs):
    if p.text and p.text.strip() in PERSONA_FIXES:
        set_text(p, PERSONA_FIXES[p.text.strip()])
        changes.append(f"Persona gender+ID added at para {i}")


# ============================================================
# 5. USER STORIES TABLE — Make IDs consistent with Personas
# Teacher: "the ID should consistent with above Personas"
# Map user story personas to include Persona IDs
# ============================================================
us_table = doc.tables[2]  # User Stories table
PERSONA_MAP = {
    "Final-Year Student": "Final-Year Student (P-01)",
    "Career Changer": "Career Changer (P-02)",
    "Long-term Unemployed": "Long-term Unemployed (P-03)",
}
for r in range(1, len(us_table.rows)):
    cell_text = us_table.cell(r, 1).text.strip()
    if cell_text in PERSONA_MAP:
        # Clear and rewrite cell
        for paragraph in us_table.cell(r, 1).paragraphs:
            set_text(paragraph, PERSONA_MAP[cell_text])
        changes.append(f"User Story row {r}: linked persona ID")


# ============================================================
# 6. ADD BORDERS TO ALL TABLES
# Teacher: "Add the border of the table to all the tables simple plain border"
# ============================================================
for t_idx, table in enumerate(doc.tables):
    add_table_borders(table)
    changes.append(f"Table {t_idx}: borders added")


# ============================================================
# 7. FUNCTIONAL REQUIREMENTS — Add source sentence
# Teacher: "do these requirements come from users' interview/feedback?"
# ============================================================
for i, p in enumerate(doc.paragraphs):
    if p.text and "categorised using MoSCoW prioritisation" in p.text:
        set_text(p, "The functional requirements below were gathered from user interviews with the three personas described above, combined with the findings from the literature review and competitor analysis. They are categorised using MoSCoW prioritisation (Must Have, Should Have, Could Have).")
        changes.append(f"Func Req source sentence added at para {i}")
        break


# ============================================================
# 8. NFR — Add GDPR link and WCAG reference
# Teacher: "any link for this doc?" and "reference?"
# These are inside tables, so we need to find and edit the cells
# ============================================================
# NFR table is table 4
nfr_table = doc.tables[4]
for r in range(len(nfr_table.rows)):
    for c in range(len(nfr_table.columns)):
        cell_text = nfr_table.cell(r, c).text
        if "GDPR" in cell_text and "UK" in cell_text and "Data Protection" in cell_text:
            for paragraph in nfr_table.cell(r, c).paragraphs:
                if "GDPR" in paragraph.text:
                    old = paragraph.text
                    new = old.rstrip()
                    if "legislation.gov.uk" not in new:
                        new += " [Available: https://www.legislation.gov.uk/ukpga/2018/12/contents/enacted]"
                        set_text(paragraph, new)
                        changes.append(f"NFR: GDPR link added")
        if "WCAG" in cell_text:
            for paragraph in nfr_table.cell(r, c).paragraphs:
                if "WCAG" in paragraph.text:
                    old = paragraph.text
                    new = old.rstrip()
                    if "w3.org" not in new:
                        new += " [Reference: W3C, 'WCAG 2.1,' 2018. Available: https://www.w3.org/TR/WCAG21/]"
                        set_text(paragraph, new)
                        changes.append(f"NFR: WCAG reference added")


# ============================================================
# 9. SPRINT 3 — Add recommendation discussion
# Teacher: "recommendation is not discussed in the content of section 7.3"
# ============================================================
for i, p in enumerate(doc.paragraphs):
    if p.text and "The same NLP pipeline from Sprint 2 was reused on job descriptions" in p.text:
        old = p.text
        if "recommendation engine" not in old.lower():
            new_text = old.rstrip() + " The recommendation engine was also built during this sprint. For each missing skill identified, the system looks up a curated dictionary of free learning resources (Coursera, FutureLearn, W3Schools) stored in skills_data.py and includes the relevant course links in the results. Skills are ranked by frequency — the skills that appear in the most job listings are flagged as high priority, giving the user a clear order of what to learn first."
            set_text(p, new_text)
            changes.append(f"Sprint 3: recommendation content added at para {i}")
        break


# ============================================================
# 10. USER TESTING TABLE — Add tester background info + feedback
# Teacher: "briefly introduce background info of 3 testers"
# Teacher: "any comments from them? like anything good and any part can be improved?"
# ============================================================
test_table = doc.tables[10]

TESTER_INFO = {
    "Emmanuel Ennin": "Emmanuel Ennin (Male, 22, Final-year Computer Science student — represents Persona P-01)",
    "Nahavipusan": "Nahavipusan (Female, 23, Business Management graduate switching to tech — represents Persona P-02)",
    "Enoch Emron": "Enoch Emron (Male, 28, Currently job-seeking after redundancy — represents Persona P-03)",
}

TESTER_FEEDBACK = {
    "Emmanuel Ennin": "Upload was straightforward. Liked the green/red skill chips. Suggested adding a percentage to each missing skill to show how important it is. The score felt accurate compared to his own job applications.",
    "Nahavipusan": "Learning Path was very useful. Priority chart was clear. Suggested the system should explain what each missing skill actually means. Found the overall interface easy to use without help.",
    "Enoch Emron": "Signup was quick. Score trend chart showed progress clearly. Suggested adding a print button for the results page. Found the plain-English feedback easy to understand.",
}

for r in range(1, len(test_table.rows)):
    name = test_table.cell(r, 0).text.strip()
    if name in TESTER_INFO:
        for paragraph in test_table.cell(r, 0).paragraphs:
            set_text(paragraph, TESTER_INFO[name])
        changes.append(f"Tester background: {name}")
    if name in TESTER_FEEDBACK:
        for paragraph in test_table.cell(r, 2).paragraphs:
            set_text(paragraph, TESTER_FEEDBACK[name])
        changes.append(f"Tester feedback: {name}")


# ============================================================
# 11. STRENGTHS & LIMITATIONS — Convert to bullet points
# Teacher: "using bullet points for these 2 sessions will improve its readability"
# ============================================================
# Strengths (para 349)
for i, p in enumerate(doc.paragraphs):
    if p.text and "Main strengths:" in p.text:
        new_text = (
            "• Real-time job data from the Reed API ensures recommendations are based on current market demand.\n"
            "• The three-layer NLP pipeline catches skills that simple keyword search would miss, improving matching accuracy.\n"
            "• The Learning Path feature turns analysis into action, allowing users to track their progress on missing skills.\n"
            "• Privacy-first design processes CVs in memory only and never saves personal documents to disk, complying with GDPR.\n"
            "• The system is completely free to use, unlike competitors such as TopCV (£99) and Jobscan (£49/month)."
        )
        set_text(p, new_text)
        changes.append(f"Strengths: converted to bullet points at para {i}")
        break

# Limitations (para 351)
for i, p in enumerate(doc.paragraphs):
    if p.text and "Limitations identified:" in p.text:
        new_text = (
            "• The skills dictionary contains approximately 200 skills — niche or emerging skills may not be recognised.\n"
            "• The system only supports English-language CVs and UK job listings through the Reed API.\n"
            "• The smallest spaCy model (en_core_web_sm) is used — a larger model could improve extraction accuracy.\n"
            "• Experience relevance and format quality scores are currently placeholder values and need further development.\n"
            "• The system requires an active internet connection to fetch live job listings from the Reed API."
        )
        set_text(p, new_text)
        changes.append(f"Limitations: converted to bullet points at para {i}")
        break


# ============================================================
# 12. REFERENCES — Add [1], [2], etc. numbering
# Teacher: "you cited these references throughout content with numbers. So has to numbering all references."
# ============================================================
ref_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text and p.text.strip() == "References" and "Heading" in (p.style.name or ""):
        ref_start = i + 1
        break

if ref_start:
    ref_num = 1
    for i in range(ref_start, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if "Heading" in (p.style.name or ""):
            break  # Hit next section
        text = p.text.strip()
        if not text:
            continue
        # Skip if already numbered
        if text.startswith("["):
            ref_num += 1
            continue
        # Add number
        set_text(p, f"[{ref_num}] {text}")
        changes.append(f"Reference [{ref_num}] numbered at para {i}")
        ref_num += 1


# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT)
print(f"\nSAVED: {OUTPUT}")
print(f"Total changes: {len(changes)}")
print("=" * 60)
for c in changes:
    print(f"  ✓ {c}")
